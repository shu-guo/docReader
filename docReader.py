import os
import docx
from docx import Document
from docx.shared import Inches
import xml.etree.ElementTree as ET
from lxml import etree
from xml.dom import minidom
import json

document = Document('demo.docx')
#print('document')
#print(document.inline_shapes)
#print(document.sections)
#print(document.part._rels)

#for paragraph in document.paragraphs:
#    print('paragraphs')

#for shape in document.inline_shapes:
#    print ('inline_shapes')

#rels = {}
#for r in document.part.rels.values():
    #if isinstance(r._target, docx.parts.image.ImagePart):
        #rels[r.rId] = os.path.basename(r._target.partname)
        #print ('rels:', rels[r.rId], r._target.partname)

tables = document.tables #获取文件中的表格集

#rIds=[]
for table in tables[:]:
    with open('export.json',"w+",encoding="utf-8") as f:
        f.write("[")
        flag=0
        for i, row in enumerate(table.rows[:]):   # 读每行
            row_content = []
            for cell in row.cells[:]:  # 读一行中的所有单元格
                for pa in cell.paragraphs:
                    p_xml_str = pa._p.xml # 按段落获取xml
                    p_xml = etree.fromstring(p_xml_str) # 转换成lxml结点
                    #print('etree:', etree.tounicode(p_xml)) # 打印查看
                    xml_dom = minidom.parseString(etree.tounicode(p_xml))
                    stus = xml_dom.getElementsByTagName('w:pict')
                    for si in stus:
                        print('si:', i, si.getElementsByTagName('v:imagedata').item(0).getAttribute('r:id'))   
                c = cell.text
                c = c.replace('\r', '\\r')
                c = c.replace('\n', '\\n')
                #print(c)
                if c.strip()=='':
                    continue
                if flag==1:
                    f.write("},\n")
                flag=1
                f.write("{\n")
                f.write('"Code":"%s"\n'% c)
                row_content.append(c)
            print (i, row_content) #以列表形式导出每一行数据
        f.write("}\n")
        f.write("]\n")


